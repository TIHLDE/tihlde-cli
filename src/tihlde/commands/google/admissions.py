import click

from InquirerPy import inquirer
from docx import Document

from tihlde.utils.google import (
    list_drive_folders,
    upload_file
)
from tihlde.utils.dir import (
    exists,
    clean_dir
)
from tihlde.settings import (
    USER_OWN_FILES,
    USER_TEMP_DIR
)
from tihlde.enums import ResponseType
from tihlde.api import (
    getGroupForms,
    getAdmissions
)
from tihlde.api.models import Form


@click.command(help="Upload admissions data to Google Cloud")
@click.option(
    "--id",
    "-i",
    type=str,
    help="Google Drive folder ID"
)
@click.option(
    "--questions",
    "-q",
    type=str,
    help="Path to the questions data file, if you want to upload questions data to the file"
)
@click.option(
    "--group",
    "-g",
    type=str,
    prompt=True,
    help="Group name"
)
def admissions(id: str, questions: str, group: str):
    """Upload admissions data to Google Cloud"""
    if not questions:
        questions = f"{USER_OWN_FILES}/admission_questions.txt"
        if not exists(questions):
            click.echo("Questions data file not found")
            return
    else:
        questions = f"{USER_OWN_FILES}/{questions}"
        if not exists(questions):
            click.echo("Questions data file not found")
            return
    
    response = getGroupForms(group)

    if response.type == ResponseType.ERROR.value:
        return click.echo(click.style(response.detail, fg="red"))
    
    forms = response.forms
    if not forms:
        return click.echo(click.style("No forms found", fg="red"))
    
    form: Form = inquirer.select(
        message="Select a form to get the admissions data from",
        choices=forms
    ).execute()

    response = getAdmissions(form.id)

    if response.type == ResponseType.ERROR.value:
        return click.echo(click.style(response.detail, fg="red"))
    
    admissions = response.admissions

    if not admissions:
        return click.echo(click.style("No admissions found", fg="red"))
    
    if not id:
        folders = list_drive_folders()
        folder = inquirer.select(
            message="Select a folder to upload the admissions data to",
            choices=folders
        ).execute()
        id = folder.id

    admission_questions: list[str] = []
    with open(questions, "r", encoding="utf-8") as file:
        for line in file:
            admission_questions.append(line.strip())

    
    with click.progressbar(admissions, label="Uploading admissions data to Google Drive") as bar:
        for admission in bar:
            doc = Document()
            file_heading = f"{admission.user.first_name} {admission.user.last_name}" 

            doc.add_heading(file_heading, level=1)
            doc.add_paragraph(f"E-post: {admission.user.email}")
            doc.add_paragraph(f"Studie: {admission.user.study.group.name}")
            doc.add_paragraph(f"Klasse: {admission.user.studyyear.group.name}")

            for field in form.fields:
                doc.add_heading(field.title)
                
                answers = admission.answers
                answer = next((answer for answer in answers if answer.field.id == field.id), None)

                if not answer:
                    doc.add_paragraph("Svar mangler")
                    continue

                if len(answer.selected_options):
                    doc.add_paragraph(", ".join(answer.selected_options))
                else:
                    doc.add_paragraph(answer.answer_text)
            
            doc.add_heading("Intervjuspørsmål", level=1)
            
            for question in admission_questions:
                doc.add_heading(question)            

            doc_path = f"{USER_TEMP_DIR}/{file_heading}.docx"
            doc.save(doc_path)
            upload_file(doc_path, id)

    clean_dir(USER_TEMP_DIR)